"""Tools fuer Terminlisten in Word-Dateien."""

from __future__ import annotations

import json
import re
import shutil
import sys
import tempfile
import threading
import time
import unicodedata
import uuid
from pathlib import Path

from docx import Document

projekt_ordner = Path(__file__).resolve().parent.parent
if str(projekt_ordner) not in sys.path:
    sys.path.insert(0, str(projekt_ordner))

from config import lade_config
from tools.audit_log import hash_file_ref, schreibe_audit_event
from tools.backup import create_backup

DATUM_MUSTER = re.compile(r"(\d{2})[._-](\d{2})[._-](\d{4})")
OEFFNEN_RETRY = 10
OEFFNEN_WARTEZEIT = 0.35
SPEICHERN_RETRY = 4
SPEICHERN_WARTEZEIT = 0.25
NACH_SPEICHERN_WARTEZEIT = 0.35
DISCOVERY_CACHE_TTL_S = 20
DOCX_INDEX_CACHE = {"expires_at": 0.0, "ordner_sig": "", "eintraege": []}
SUCHINDEX_DATEI = projekt_ordner / "agent" / "word_search_index.json"
SUCHINDEX_CACHE_TTL_S = 25
SUCHINDEX_FORMAT = 1
SUCHINDEX_CACHE = {"expires_at": 0.0, "ordner_sig": "", "daten": {}}
SUCHINDEX_LOCK = threading.Lock()
PENDING_WRITES_DATEI = projekt_ordner / "logs" / "pending_writes.json"
PENDING_TIMEOUT_S = 600
PENDING_SLEEP_S = 4
PENDING_LOCK = threading.Lock()
PENDING_WORKER_GESTARTET = False


def _fehler_result(text: str, code: str) -> dict:
    return {"fehler": str(text or ""), "fehler_code": str(code or "")}


def _pending_fingerprint(filepath: str) -> dict:
    mtime_ns, groesse = _datei_stat(Path(str(filepath or "")))
    return {"mtime_ns": int(mtime_ns), "groesse": int(groesse)}


def _pending_laden() -> list[dict]:
    try:
        if not PENDING_WRITES_DATEI.exists():
            return []
        daten = json.loads(PENDING_WRITES_DATEI.read_text(encoding="utf-8"))
        if isinstance(daten, list):
            return [eintrag for eintrag in daten if isinstance(eintrag, dict)]
    except Exception:
        pass
    return []


def _pending_speichern(eintraege: list[dict]) -> None:
    try:
        PENDING_WRITES_DATEI.parent.mkdir(parents=True, exist_ok=True)
        payload = [eintrag for eintrag in eintraege if isinstance(eintrag, dict)]
        PENDING_WRITES_DATEI.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def _pending_eintrag(filepath: str, uhrzeit: str, werte: dict, ersetzen: bool, backup_pfad: str) -> dict:
    jetzt = int(time.time())
    return {
        "operation_id": uuid.uuid4().hex,
        "filepath": str(filepath or "").strip(),
        "uhrzeit": str(uhrzeit or "").strip(),
        "werte": dict(werte or {}),
        "ersetzen": bool(ersetzen),
        "backup_pfad": str(backup_pfad or "").strip(),
        "queued_at": jetzt,
        "expires_at": jetzt + int(PENDING_TIMEOUT_S),
        "next_retry_at": jetzt,
        "retry_count": 0,
        "status": "pending",
        "fingerprint": _pending_fingerprint(filepath),
    }


def _pending_requeue(eintrag: dict) -> dict:
    retry = int(eintrag.get("retry_count", 0) or 0) + 1
    delay = min(45, 2 + retry * 2)
    neu = dict(eintrag)
    neu["retry_count"] = retry
    neu["next_retry_at"] = int(time.time()) + delay
    neu["status"] = "retrying"
    return neu


def _pending_due(eintrag: dict, jetzt: int) -> bool:
    if int(eintrag.get("expires_at", 0) or 0) <= jetzt:
        return True
    return int(eintrag.get("next_retry_at", 0) or 0) <= jetzt


def _pending_ablauf_hinweis() -> str:
    return "Datei war zu lang gesperrt – bitte Word schließen und erneut bestätigen."


def _pending_driftschutz(eintrag: dict) -> bool:
    alt = eintrag.get("fingerprint", {}) if isinstance(eintrag, dict) else {}
    neu = _pending_fingerprint(str(eintrag.get("filepath", "") or ""))
    return int(alt.get("mtime_ns", -1)) == int(neu.get("mtime_ns", -2)) and int(alt.get("groesse", -1)) == int(neu.get("groesse", -2))


def _pending_apply(dokument, eintrag: dict) -> dict:
    if not dokument.tables:
        return _fehler_result("Keine Tabelle in der Datei gefunden", "no_table")
    zeile = _finde_zeile_fuer_uhrzeit(dokument.tables[0], str(eintrag.get("uhrzeit", "") or ""))
    if zeile is None:
        return _fehler_result("Slot nicht gefunden", "slot_not_found")
    ist_belegt = _slot_ist_belegt(zeile)
    if ist_belegt and not bool(eintrag.get("ersetzen", False)):
        return _fehler_result("Slot bereits belegt", "slot_belegt")
    neu = dict(eintrag.get("werte", {}) if isinstance(eintrag.get("werte"), dict) else {})
    alt = _lese_terminwerte_aus_zeile(zeile) if ist_belegt else {}
    ziel = _merge_terminwerte(neu, alt, bool(eintrag.get("ersetzen", False)))
    if not str(ziel.get("name", "") or "").strip():
        return _fehler_result("Name fehlt", "name_required")
    _trage_terminwerte_ein(zeile, ziel)
    return {"erfolg": True}


def _pending_verarbeite_eintrag(eintrag: dict, ordner_liste: list[Path], jetzt: int) -> tuple[dict | None, dict | None]:
    if int(eintrag.get("expires_at", 0) or 0) <= jetzt:
        _debug_word_ereignis("word.write.expired", "write_expired", str(eintrag.get("filepath", "")))
        abgelaufen = dict(eintrag)
        abgelaufen["status"] = "expired"
        abgelaufen["expired_message"] = _pending_ablauf_hinweis()
        abgelaufen["next_retry_at"] = int(zeit := time.time()) + 3600
        abgelaufen["expires_at"] = int(zeit) + 3600
        return abgelaufen, _fehler_result(_pending_ablauf_hinweis(), "expired")
    if not _pending_driftschutz(eintrag):
        _debug_word_ereignis("word.write.pending", "manual_action_required", str(eintrag.get("filepath", "")))
        return None, _fehler_result("Datei hat sich geaendert. Bitte erneut bestaetigen.", "manual_action_required")
    dokument, fehler = _oeffne_dokument(str(eintrag.get("filepath", "")), ordner_liste, read_only=False)
    if fehler:
        _debug_word_ereignis("word.write.retry", "lock_detected", str(eintrag.get("filepath", "")), extra_meta={"retry": int(eintrag.get("retry_count", 0) or 0)})
        return _pending_requeue(eintrag), None
    apply_result = _pending_apply(dokument, eintrag)
    if apply_result.get("fehler"):
        return None, apply_result
    save_fehler = _speichere_dokument(dokument, str(eintrag.get("filepath", "")))
    if save_fehler:
        return _pending_requeue(eintrag), None
    _debug_word_ereignis("word.write.committed", "success_after_retry", str(eintrag.get("filepath", "")))
    return None, {"erfolg": True, "status": "success_after_retry", "backup_pfad": str(eintrag.get("backup_pfad", "") or "")}


def _pending_worker_tick() -> None:
    ordner_liste, fehler = _hole_erlaubte_ordner()
    if fehler:
        return
    with PENDING_LOCK:
        jetzt = int(time.time())
        ausgang = _pending_laden()
        neu = []
        for eintrag in ausgang:
            if str(eintrag.get("status", "")) == "expired":
                neu.append(eintrag)
                continue
            if not _pending_due(eintrag, jetzt):
                neu.append(eintrag)
                continue
            aktualisiert, _ = _pending_verarbeite_eintrag(eintrag, ordner_liste, jetzt)
            if isinstance(aktualisiert, dict):
                neu.append(aktualisiert)
        _pending_speichern(neu)


def _pending_expired_hinweis(filepath: str, uhrzeit: str) -> str:
    with PENDING_LOCK:
        daten = _pending_laden()
        rest = []
        hinweis = ""
        for eintrag in daten:
            gleich_datei = str(eintrag.get("filepath", "") or "").strip() == str(filepath or "").strip()
            gleich_uhrzeit = str(eintrag.get("uhrzeit", "") or "").strip() == str(uhrzeit or "").strip()
            ist_expired = str(eintrag.get("status", "") or "") == "expired"
            if gleich_datei and gleich_uhrzeit and ist_expired and not hinweis:
                hinweis = str(eintrag.get("expired_message", "") or _pending_ablauf_hinweis())
                continue
            rest.append(eintrag)
        if len(rest) != len(daten):
            _pending_speichern(rest)
        return hinweis


def _pending_worker_loop() -> None:
    while True:
        try:
            _pending_worker_tick()
        except Exception:
            pass
        time.sleep(PENDING_SLEEP_S)


def _pending_worker_starten() -> None:
    global PENDING_WORKER_GESTARTET
    if PENDING_WORKER_GESTARTET:
        return
    PENDING_WORKER_GESTARTET = True
    threading.Thread(target=_pending_worker_loop, name="pending-write-worker", daemon=True).start()


def _pending_enqueue(filepath: str, uhrzeit: str, werte: dict, ersetzen: bool, backup_pfad: str) -> dict:
    eintrag = _pending_eintrag(filepath, uhrzeit, werte, ersetzen, backup_pfad)
    with PENDING_LOCK:
        daten = _pending_laden()
        daten.append(eintrag)
        _pending_speichern(daten)
    _debug_word_ereignis("word.write.pending", "pending_enqueued", filepath, extra_meta={"operation_id": eintrag.get("operation_id", "")})
    _pending_worker_starten()
    return {
        "erfolg": False,
        "status": "pending",
        "pending_id": str(eintrag.get("operation_id", "") or ""),
        "backup_pfad": backup_pfad,
        "fehler": "Datei ist gesperrt. Schreibauftrag wurde in die Warteschlange gelegt.",
        "fehler_code": "pending_write",
    }


def _ist_lock_fehler(fehlertext: str) -> bool:
    text = str(fehlertext or "").strip().lower()
    return "gesperrt" in text or "lock" in text


def _debug_word_ereignis(
    aktion: str,
    code: str = "",
    filepath: str = "",
    detail: str = "",
    extra_meta: dict | None = None,
) -> None:
    meta = {"tool_name": "word_tools"}
    pfad = str(filepath or "").strip()
    if pfad:
        meta["file_ref_hash"] = hash_file_ref(pfad)
        mtime_ns, groesse = _datei_stat(Path(pfad))
        if mtime_ns >= 0:
            meta["file_mtime_ns"] = int(mtime_ns)
        if groesse >= 0:
            meta["file_size"] = int(groesse)
    if code:
        meta["tool_error_code"] = str(code)[:120]
    if detail:
        meta["block_reason"] = str(detail)[:120]
    if isinstance(extra_meta, dict):
        for key, value in extra_meta.items():
            k = str(key or "").strip()
            if not k:
                continue
            if isinstance(value, (str, int, float, bool)):
                meta[k] = value
            else:
                meta[k] = str(value)
    try:
        schreibe_audit_event(
            action=str(aktion or "word.debug"),
            result="error" if code else "ok",
            error_code=str(code or ""),
            meta=meta,
        )
    except Exception:
        pass


def _normalisiere_datum(datum: str) -> str:
    text = str(datum or "").strip()
    treffer = re.search(r"(\d{1,2})\s*[.\-_/ ]\s*(\d{1,2})\s*[.\-_/ ]\s*(\d{4})", text)
    if treffer is None:
        return ""
    tag, monat, jahr = treffer.groups()
    if not (1 <= int(tag) <= 31 and 1 <= int(monat) <= 12):
        return ""
    return f"{int(tag):02d}_{int(monat):02d}_{jahr}"


def _datum_aus_dateiname(dateiname: str) -> str:
    treffer = DATUM_MUSTER.search(dateiname)
    if treffer is None:
        return ""
    tag, monat, jahr = treffer.groups()
    return f"{tag}_{monat}_{jahr}"


def _datum_varianten(datum_normalisiert: str) -> set[str]:
    datum = str(datum_normalisiert or "").strip()
    if not datum:
        return set()
    teile = datum.split("_")
    if len(teile) != 3:
        return {datum}
    tag, monat, jahr = teile
    return {
        f"{tag}_{monat}_{jahr}",
        f"{tag}.{monat}.{jahr}",
        f"{tag}-{monat}-{jahr}",
    }


def _ordner_signatur(ordner_liste: list[Path]) -> str:
    teile = [str(ordner) for ordner in ordner_liste]
    return "|".join(sorted(teile))


def _datei_stat(datei: Path) -> tuple[int, int]:
    try:
        stat = datei.stat()
        return int(stat.st_mtime_ns), int(stat.st_size)
    except Exception:
        return -1, -1


def _index_eintrag(datei: Path) -> dict:
    mtime_ns, groesse = _datei_stat(datei)
    return {
        "dateiname": datei.name,
        "pfad": str(datei),
        "datum": _datum_aus_dateiname(datei.name),
        "mtime_ns": mtime_ns,
        "groesse": groesse,
    }


def _scanne_docx_index(ordner_liste: list[Path]) -> list[dict]:
    eintraege = []
    for ordner in ordner_liste:
        for datei in ordner.rglob("*.docx"):
            if not _ist_in_erlaubten_ordnern(datei, ordner_liste):
                continue
            eintraege.append(_index_eintrag(datei))
    return sorted(eintraege, key=lambda eintrag: eintrag["dateiname"].lower())


def _cache_eintraege_gueltig(eintraege: list[dict]) -> bool:
    for eintrag in eintraege:
        pfad = Path(str(eintrag.get("pfad", "") or ""))
        if not pfad.exists():
            return False
        mtime_ns, groesse = _datei_stat(pfad)
        if mtime_ns != int(eintrag.get("mtime_ns", -2)):
            return False
        if groesse != int(eintrag.get("groesse", -2)):
            return False
    return True


def _cache_ist_gueltig(ordner_sig: str) -> bool:
    ablauf = float(DOCX_INDEX_CACHE.get("expires_at", 0.0) or 0.0)
    cache_sig = str(DOCX_INDEX_CACHE.get("ordner_sig", "") or "")
    if time.time() > ablauf or cache_sig != ordner_sig:
        return False
    eintraege = DOCX_INDEX_CACHE.get("eintraege", [])
    if not isinstance(eintraege, list):
        return False
    return _cache_eintraege_gueltig(eintraege)


def _hole_docx_index(ordner_liste: list[Path]) -> list[dict]:
    global DOCX_INDEX_CACHE
    ordner_sig = _ordner_signatur(ordner_liste)
    if _cache_ist_gueltig(ordner_sig):
        eintraege = DOCX_INDEX_CACHE.get("eintraege", [])
        return list(eintraege) if isinstance(eintraege, list) else []
    eintraege = _scanne_docx_index(ordner_liste)
    DOCX_INDEX_CACHE = {
        "expires_at": time.time() + float(DISCOVERY_CACHE_TTL_S),
        "ordner_sig": ordner_sig,
        "eintraege": list(eintraege),
    }
    return eintraege


def _index_feature_aktiv() -> bool:
    config = lade_config()
    flags = config.get("feature_flags", {})
    if not isinstance(flags, dict):
        return True
    return bool(flags.get("word_search_index", True))


def _lade_suchindex_von_disk(ordner_sig: str) -> dict:
    try:
        if not SUCHINDEX_DATEI.exists():
            return {}
        daten = json.loads(SUCHINDEX_DATEI.read_text(encoding="utf-8"))
        if int(daten.get("format", 0)) != SUCHINDEX_FORMAT:
            return {}
        if str(daten.get("ordner_sig", "")) != str(ordner_sig):
            return {}
        inhalt = daten.get("daten", {})
        return inhalt if isinstance(inhalt, dict) else {}
    except Exception:
        return {}


def _speichere_suchindex_auf_disk(ordner_sig: str, daten: dict) -> None:
    payload = {"format": SUCHINDEX_FORMAT, "ordner_sig": ordner_sig, "daten": daten}
    try:
        SUCHINDEX_DATEI.parent.mkdir(parents=True, exist_ok=True)
        SUCHINDEX_DATEI.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _cache_suchindex_gueltig(ordner_sig: str) -> bool:
    if time.time() > float(SUCHINDEX_CACHE.get("expires_at", 0.0) or 0.0):
        return False
    if str(SUCHINDEX_CACHE.get("ordner_sig", "")) != str(ordner_sig):
        return False
    return isinstance(SUCHINDEX_CACHE.get("daten", {}), dict)


def _uebernehme_suchindex_in_cache(ordner_sig: str, daten: dict) -> dict:
    global SUCHINDEX_CACHE
    SUCHINDEX_CACHE = {
        "expires_at": time.time() + float(SUCHINDEX_CACHE_TTL_S),
        "ordner_sig": str(ordner_sig),
        "daten": dict(daten or {}),
    }
    return dict(SUCHINDEX_CACHE.get("daten", {}))


def _lese_suchindex(ordner_sig: str) -> dict:
    if _cache_suchindex_gueltig(ordner_sig):
        daten = SUCHINDEX_CACHE.get("daten", {})
        return dict(daten) if isinstance(daten, dict) else {}
    daten = _lade_suchindex_von_disk(ordner_sig)
    return _uebernehme_suchindex_in_cache(ordner_sig, daten)


def _eintrag_unveraendert(alt: dict, mtime_ns: int, groesse: int, datum: str) -> bool:
    if int(alt.get("mtime_ns", -1)) != int(mtime_ns):
        return False
    if int(alt.get("groesse", -1)) != int(groesse):
        return False
    if str(alt.get("datum", "")) != str(datum):
        return False
    return isinstance(alt.get("treffer", []), list)


def _zeilen_treffer_aus_dokument(dokument, filepath: str, datei_datum: str) -> list[dict]:
    if not dokument.tables:
        return []
    treffer = []
    for zeile in dokument.tables[0].rows[1:]:
        daten = _eintrag_aus_zeile(zeile, filepath, datei_datum)
        if not daten:
            continue
        zeilen_text = " ".join(str(z.text or "") for z in zeile.cells)
        treffer.append({
            "uhrzeit": str(daten.get("uhrzeit", "") or ""),
            "name": str(daten.get("name", "") or ""),
            "such_norm": _text_normalisieren(zeilen_text),
        })
    return treffer


def _baue_suchindex_fuer_datei(dateipfad: str, datei_datum: str, ordner_liste: list[Path]) -> tuple[dict, str]:
    dokument, fehler = _lese_dokument_fuer_suche(dateipfad, ordner_liste)
    if fehler and fehler != "no_table":
        return {}, fehler
    mtime_ns, groesse = _datei_stat(Path(dateipfad))
    treffer = _zeilen_treffer_aus_dokument(dokument, dateipfad, datei_datum) if dokument else []
    return {
        "mtime_ns": int(mtime_ns),
        "groesse": int(groesse),
        "datum": str(datei_datum or ""),
        "treffer": treffer,
    }, ""


def _aktualisiere_suchindex(dateien: list[dict], ordner_liste: list[Path], basis: dict) -> tuple[dict, int]:
    neue_daten = {}
    fehler_anzahl = 0
    for eintrag in dateien:
        dateipfad = str(eintrag.get("pfad", "") or "").strip()
        datei_datum = str(eintrag.get("datum", "") or "")
        mtime_ns, groesse = _datei_stat(Path(dateipfad))
        alt = basis.get(dateipfad, {}) if isinstance(basis, dict) else {}
        if _eintrag_unveraendert(alt, mtime_ns, groesse, datei_datum):
            neue_daten[dateipfad] = dict(alt)
            continue
        neu, fehler = _baue_suchindex_fuer_datei(dateipfad, datei_datum, ordner_liste)
        if fehler:
            fehler_anzahl += 1
            continue
        neue_daten[dateipfad] = neu
    return neue_daten, fehler_anzahl


def _suche_im_suchindex(dateien: list[dict], such_norm: str, zielzeit: str, daten: dict) -> list[dict]:
    treffer = []
    for eintrag in dateien:
        dateipfad = str(eintrag.get("pfad", "") or "").strip()
        dateiname = str(eintrag.get("dateiname", "") or Path(dateipfad).name)
        datei_datum = str(eintrag.get("datum", "") or "")
        datei_index = daten.get(dateipfad, {}) if isinstance(daten, dict) else {}
        for kandidat in datei_index.get("treffer", []):
            if zielzeit and str(kandidat.get("uhrzeit", "")) != zielzeit:
                continue
            if such_norm and such_norm not in str(kandidat.get("such_norm", "")):
                continue
            treffer.append({
                "datum": datei_datum,
                "dateiname": dateiname,
                "filepath": dateipfad,
                "uhrzeit": str(kandidat.get("uhrzeit", "") or ""),
                "name": str(kandidat.get("name", "") or ""),
            })
    return treffer


def _logge_such_performance(quelle: str, dauer_ms: int, treffer: int, dateien: int) -> None:
    meta = {
        "tool_name": "word_tools",
        "quelle": str(quelle or "scan"),
        "treffer": int(treffer),
        "roh_treffer": int(dateien),
        "block_reason": f"search_ms={max(0, int(dauer_ms))}",
    }
    try:
        schreibe_audit_event(action="word.search.performance", duration_ms=max(0, int(dauer_ms)), meta=meta)
    except Exception:
        pass


def _logge_such_performance_erweitert(
    quelle: str,
    dauer_ms: int,
    treffer: int,
    dateien: int,
    stats: dict | None = None,
) -> None:
    meta = {
        "tool_name": "word_tools",
        "quelle": str(quelle or "scan"),
        "treffer": int(treffer),
        "roh_treffer": int(dateien),
        "block_reason": f"search_ms={max(0, int(dauer_ms))}",
    }
    if isinstance(stats, dict):
        for key in [
            "opened",
            "skipped",
            "retry_opened",
            "retry_skipped",
            "limit_hit",
            "max_open_attempts",
            "max_skip",
            "stages",
        ]:
            if key in stats:
                meta[key] = stats.get(key)
        if stats.get("limit_reason"):
            meta["limit_reason"] = str(stats.get("limit_reason"))
    try:
        schreibe_audit_event(action="word.search.performance", duration_ms=max(0, int(dauer_ms)), meta=meta)
    except Exception:
        pass


def _scan_guardrail_limits(stufe: str) -> dict:
    name = str(stufe or "").strip().lower()
    if name == "full":
        return {"max_open_attempts": 320, "max_skip": 160}
    return {"max_open_attempts": 80, "max_skip": 45}


def _scan_stats_neu(max_open_attempts: int, max_skip: int, stufe: str) -> dict:
    return {
        "opened": 0,
        "skipped": 0,
        "retry_opened": 0,
        "retry_skipped": 0,
        "limit_hit": False,
        "limit_reason": "",
        "max_open_attempts": max(0, int(max_open_attempts or 0)),
        "max_skip": max(0, int(max_skip or 0)),
        "stages": str(stufe or "scan"),
    }


def _scan_versuche(stats: dict) -> int:
    return int(stats.get("opened", 0)) + int(stats.get("skipped", 0))


def _scan_limit_erreicht(stats: dict) -> bool:
    max_open = int(stats.get("max_open_attempts", 0) or 0)
    max_skip = int(stats.get("max_skip", 0) or 0)
    if max_open > 0 and _scan_versuche(stats) >= max_open:
        stats["limit_hit"] = True
        stats["limit_reason"] = "max_open_attempts"
        return True
    if max_skip > 0 and int(stats.get("skipped", 0)) >= max_skip:
        stats["limit_hit"] = True
        stats["limit_reason"] = "max_skip"
        return True
    return False


def _search_appointments_scan(
    datei_suche: dict,
    ordner_liste: list[Path],
    such_norm: str,
    zielzeit: str,
    max_open_attempts: int,
    max_skip: int,
    stufe: str,
) -> dict:
    treffer = []
    fehlerpfade = []
    daten_map = {}
    stats = _scan_stats_neu(max_open_attempts, max_skip, stufe)
    for eintrag in datei_suche.get("dateien", []):
        if _scan_limit_erreicht(stats):
            break
        dateipfad = str(eintrag.get("pfad", "") or "").strip()
        daten_map[dateipfad] = str(eintrag.get("datum", "") or "")
        dokument, lesefehler = _lese_dokument_fuer_suche(dateipfad, ordner_liste)
        if lesefehler:
            fehlerpfade.append(dateipfad)
            stats["skipped"] = int(stats.get("skipped", 0)) + 1
            _scan_limit_erreicht(stats)
            continue
        stats["opened"] = int(stats.get("opened", 0)) + 1
        datei_datum = str(eintrag.get("datum", "") or "")
        _fuege_treffer_aus_dokument_hinzu(dokument, dateipfad, datei_datum, such_norm, zielzeit, treffer)
    noch_fehler = _retry_gescheiterte_dateien(fehlerpfade, ordner_liste, treffer, such_norm, zielzeit, daten_map, stats)
    if _scan_limit_erreicht(stats):
        warnung = "Suche wurde aus Performance-Gruenden frueh beendet"
        return {
            "treffer": treffer,
            "warnung": warnung,
            "erweiterter_vollscan_moeglich": True,
            "scan_stats": stats,
        }
    if noch_fehler:
        return {"treffer": treffer, "warnung": "Einige Dateien konnten nicht gelesen werden", "scan_stats": stats}
    return {"treffer": treffer, "scan_stats": stats}


def _search_appointments_index(datei_suche: dict, ordner_liste: list[Path], such_norm: str, zielzeit: str) -> dict:
    try:
        dateien = datei_suche.get("dateien", [])
        if not isinstance(dateien, list):
            return {"fehler": "index_input"}
        ordner_sig = _ordner_signatur(ordner_liste)
        with SUCHINDEX_LOCK:
            basis = _lese_suchindex(ordner_sig)
            daten, fehler = _aktualisiere_suchindex(dateien, ordner_liste, basis)
            _uebernehme_suchindex_in_cache(ordner_sig, daten)
            _speichere_suchindex_auf_disk(ordner_sig, daten)
        treffer = _suche_im_suchindex(dateien, such_norm, zielzeit, daten)
        if fehler:
            return {"treffer": treffer, "warnung": "Einige Dateien konnten nicht gelesen werden"}
        return {"treffer": treffer}
    except Exception:
        return {"fehler": "index_exception"}


def warmup_search_index() -> dict:
    if not _index_feature_aktiv():
        return {"erfolg": True, "quelle": "deaktiviert"}
    ordner_liste, fehler = _hole_erlaubte_ordner()
    if fehler:
        return {"fehler": fehler}
    datei_suche = {"dateien": _hole_docx_index(ordner_liste)}
    suchergebnis = _search_appointments_index(datei_suche, ordner_liste, "", "")
    if "fehler" in suchergebnis:
        return {"fehler": str(suchergebnis.get("fehler", "index_fehler"))}
    return {"erfolg": True, "dateien": len(datei_suche.get("dateien", []))}


def _wartezeit_oeffnen(versuch: int) -> float:
    basis = float(OEFFNEN_WARTEZEIT)
    exponent = max(0, int(versuch or 0))
    return min(4.0, basis * (2 ** exponent))


def _ordnerliste_aus_config(config: dict) -> list[str]:
    pfade = config.get("pfade", {})
    ordner = pfade.get("erlaubte_ordner", [])
    if isinstance(ordner, list):
        return [str(eintrag).strip() for eintrag in ordner if str(eintrag).strip()]
    alt = str(pfade.get("terminlisten_ordner", "")).strip()
    if alt:
        return [alt]
    return []


def _hole_erlaubte_ordner() -> tuple[list[Path], str]:
    config = lade_config()
    ordner_roh = _ordnerliste_aus_config(config)
    if not ordner_roh:
        return [], "Keine erlaubten Ordner konfiguriert"
    ordner = [Path(eintrag).resolve() for eintrag in ordner_roh]
    for kandidat in ordner:
        if not kandidat.exists() or not kandidat.is_dir():
            return [], f"Ordner existiert nicht: {kandidat}"
    return ordner, ""


def _ist_in_erlaubten_ordnern(datei: Path, erlaubte_ordner: list[Path]) -> bool:
    try:
        datei_resolved = datei.resolve()
        for ordner in erlaubte_ordner:
            datei_resolved.relative_to(ordner)
            return True
    except ValueError:
        return False
    return False


def _oeffne_docx_mit_retry(datei: Path):
    letzter_fehler = "Dokument konnte nicht geoeffnet werden"
    letzter_fehler_typ = ""
    letzter_versuch = 0
    for versuch in range(OEFFNEN_RETRY):
        letzter_versuch = versuch + 1
        try:
            return Document(str(datei)), ""
        except PermissionError as exc:
            letzter_fehler = "Datei ist kurzfristig gesperrt"
            letzter_fehler_typ = exc.__class__.__name__
        except OSError as exc:
            letzter_fehler = "Datei konnte nicht geoeffnet werden"
            letzter_fehler_typ = exc.__class__.__name__
        except Exception as exc:
            letzter_fehler = "Dokument konnte nicht geoeffnet werden"
            letzter_fehler_typ = exc.__class__.__name__
        if versuch < OEFFNEN_RETRY - 1:
            time.sleep(_wartezeit_oeffnen(versuch))
    _debug_word_ereignis(
        "word.open.failed",
        "open_retry_failed",
        str(datei),
        f"retry={OEFFNEN_RETRY};fehler={letzter_fehler}",
        {
            "retry_last": letzter_versuch,
            "exception_type": letzter_fehler_typ,
            "fallback_used": False,
        },
    )
    return None, letzter_fehler


def _temp_copy_fallback_aktiv() -> bool:
    config = lade_config()
    flags = config.get("feature_flags", {})
    if not isinstance(flags, dict):
        return True
    return bool(flags.get("word_temp_copy_fallback", True))


def _oeffne_docx_ueber_tempkopie(datei: Path):
    temp_pfad = ""
    try:
        with tempfile.NamedTemporaryFile(prefix="praxis_docx_", suffix=".docx", delete=False) as temp_datei:
            temp_pfad = temp_datei.name
        shutil.copy2(str(datei), temp_pfad)
        dokument = Document(temp_pfad)
        _debug_word_ereignis(
            "word.open.fallback",
            "temp_copy_success",
            str(datei),
            "opened_via_temp_copy",
            {"fallback_used": True},
        )
        return dokument, ""
    except Exception as exc:
        _debug_word_ereignis(
            "word.open.fallback",
            "temp_copy_failed",
            str(datei),
            "fallback_failed",
            {"exception_type": exc.__class__.__name__, "fallback_used": True},
        )
        return None, "Dokument konnte nicht geoeffnet werden"
    finally:
        if temp_pfad:
            try:
                Path(temp_pfad).unlink(missing_ok=True)
            except Exception:
                pass


def _logge_read_fallback_stage(stage: str, filepath: str, erfolg: bool, detail: str = "") -> None:
    pfad = str(filepath or "").strip()
    meta = {"tool_name": "word_tools", "stage": str(stage or "")}
    if pfad:
        meta["file_ref_hash"] = hash_file_ref(pfad)
    if detail:
        meta["block_reason"] = str(detail)[:120]
    try:
        schreibe_audit_event(
            action="word.read.fallback.stage",
            result="ok" if bool(erfolg) else "error",
            error_code="" if bool(erfolg) else "stage_failed",
            meta=meta,
        )
    except Exception:
        pass


def _logge_read_fallback_stage_detail(stage: str, filepath: str, erfolg: bool, detail: str, extra_meta: dict) -> None:
    pfad = str(filepath or "").strip()
    meta = {"tool_name": "word_tools", "stage": str(stage or "")}
    if pfad:
        meta["file_ref_hash"] = hash_file_ref(pfad)
    if detail:
        meta["block_reason"] = str(detail)[:120]
    for key, wert in (extra_meta or {}).items():
        if isinstance(wert, (str, int, float, bool)):
            meta[str(key)] = wert
    try:
        schreibe_audit_event(
            action="word.read.fallback.stage",
            result="ok" if bool(erfolg) else "error",
            error_code="" if bool(erfolg) else "stage_failed",
            meta=meta,
        )
    except Exception:
        pass


def _ausnahme_detail(exc: Exception) -> str:
    typ = exc.__class__.__name__
    text = str(exc or "").strip().replace("\n", " ")
    return f"{typ}:{text}"[:120] if text else typ[:120]


def _normalisiere_pfad_schluessel(pfad: str) -> str:
    try:
        return str(Path(str(pfad or "")).resolve()).replace("/", "\\").lower()
    except Exception:
        return str(pfad or "").replace("/", "\\").lower()


def _pfad_varianten(pfad: str) -> set[str]:
    text = str(pfad or "").strip()
    if not text:
        return set()
    varianten = {text, text.replace("/", "\\")}
    try:
        path_obj = Path(text)
        varianten.add(str(path_obj))
        varianten.add(str(path_obj.resolve(strict=False)))
    except Exception:
        pass
    normalisiert = set()
    for kandidat in varianten:
        normalisiert.add(_normalisiere_pfad_schluessel(kandidat))
    return {eintrag for eintrag in normalisiert if eintrag}


def _gleicher_pfad(a: str, b: str) -> bool:
    va = _pfad_varianten(a)
    vb = _pfad_varianten(b)
    return bool(va and vb and va.intersection(vb))


def _finde_geoeffnetes_word_dokument(app, ziel_datei: Path):
    ziel_pfad = str(ziel_datei)
    ziel_name = str(ziel_datei.name).lower()
    name_treffer = []
    for index in range(1, int(app.Documents.Count) + 1):
        dokument = app.Documents.Item(index)
        vollname = str(getattr(dokument, "FullName", "") or "")
        if _gleicher_pfad(vollname, ziel_pfad):
            return dokument, "path_match"
        if str(getattr(dokument, "Name", "") or "").lower() == ziel_name:
            name_treffer.append(dokument)
    if len(name_treffer) == 1:
        return name_treffer[0], "name_match_unique"
    detail = f"not_found_open_docs={int(app.Documents.Count)}"
    if name_treffer:
        detail = f"name_match_ambiguous={len(name_treffer)}"
    return None, detail


def _lade_docx_aus_word_kopie(dokument, prefix: str):
    temp_pfad = ""
    try:
        temp_pfad = str(Path(tempfile.gettempdir()) / f"{prefix}{uuid.uuid4().hex}.docx")
        try:
            dokument.SaveCopyAs2(temp_pfad, FileFormat=16)
        except Exception:
            dokument.SaveCopyAs(temp_pfad)
        return Document(temp_pfad), ""
    except Exception as exc:
        return None, _ausnahme_detail(exc)
    finally:
        if temp_pfad:
            try:
                Path(temp_pfad).unlink(missing_ok=True)
            except Exception:
                pass


def _oeffne_docx_ueber_laufendes_word(datei: Path):
    com_aktiv = False
    try:
        import pythoncom  # type: ignore
        from win32com.client import GetActiveObject  # type: ignore
        pythoncom.CoInitialize()
        com_aktiv = True
        app = GetActiveObject("Word.Application")
        dokument, detail = _finde_geoeffnetes_word_dokument(app, datei)
        if dokument is not None:
            docx, fehler = _lade_docx_aus_word_kopie(dokument, "praxis_docx_live_")
            if not fehler:
                _logge_read_fallback_stage_detail(
                    "stage_b_running_word",
                    str(datei),
                    True,
                    "save_copy_from_open_doc",
                    {"match_detail": detail},
                )
                return docx, ""
            _logge_read_fallback_stage_detail(
                "stage_b_running_word",
                str(datei),
                False,
                "save_copy_failed",
                {"match_detail": detail, "exception": str(fehler or "")[:120]},
            )
            return None, "Dokument konnte nicht geoeffnet werden"
        _logge_read_fallback_stage_detail(
            "stage_b_running_word",
            str(datei),
            False,
            "not_found_in_running_instance",
            {"match_detail": detail},
        )
        return None, "Dokument konnte nicht geoeffnet werden"
    except Exception as exc:
        _logge_read_fallback_stage_detail(
            "stage_b_running_word",
            str(datei),
            False,
            "running_word_error",
            {"exception": _ausnahme_detail(exc)},
        )
        return None, "Dokument konnte nicht geoeffnet werden"
    finally:
        if com_aktiv:
            try:
                pythoncom.CoUninitialize()  # type: ignore[name-defined]
            except Exception:
                pass


def _oeffne_docx_ueber_word_getobject(datei: Path):
    com_aktiv = False
    try:
        import pythoncom  # type: ignore
        from win32com.client import GetObject  # type: ignore
        pythoncom.CoInitialize()
        com_aktiv = True
        dokument = GetObject(str(datei))
        docx, fehler = _lade_docx_aus_word_kopie(dokument, "praxis_docx_getobj_")
        if fehler:
            _logge_read_fallback_stage_detail(
                "stage_b2_getobject",
                str(datei),
                False,
                "save_copy_failed",
                {},
            )
            return None, "Dokument konnte nicht geoeffnet werden"
        _logge_read_fallback_stage_detail(
            "stage_b2_getobject",
            str(datei),
            True,
            "save_copy_from_getobject",
            {},
        )
        return docx, ""
    except Exception as exc:
        _logge_read_fallback_stage_detail(
            "stage_b2_getobject",
            str(datei),
            False,
            "getobject_error",
            {"exception": _ausnahme_detail(exc)},
        )
        return None, "Dokument konnte nicht geoeffnet werden"
    finally:
        if com_aktiv:
            try:
                pythoncom.CoUninitialize()  # type: ignore[name-defined]
            except Exception:
                pass


def _oeffne_docx_ueber_com_readonly(datei: Path):
    com_aktiv = False
    temp_pfad = ""
    app = None
    dokument = None
    try:
        import pythoncom  # type: ignore
        from win32com.client import DispatchEx  # type: ignore
        pythoncom.CoInitialize()
        com_aktiv = True
        app = DispatchEx("Word.Application")
        app.Visible = False
        app.DisplayAlerts = 0
        try:
            app.AutomationSecurity = 3
        except Exception:
            pass
        dokument = app.Documents.Open(
            FileName=str(datei),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
            OpenAndRepair=True,
        )
        temp_pfad = str(Path(tempfile.gettempdir()) / f"praxis_docx_com_{uuid.uuid4().hex}.docx")
        try:
            dokument.SaveAs2(temp_pfad, FileFormat=16)
        except Exception:
            dokument.SaveCopyAs(temp_pfad)
        docx = Document(temp_pfad)
        _logge_read_fallback_stage("stage_c_com_readonly", str(datei), True, "com_readonly_copy_success")
        return docx, ""
    except Exception as exc:
        _logge_read_fallback_stage_detail(
            "stage_c_com_readonly",
            str(datei),
            False,
            "com_error",
            {"exception": _ausnahme_detail(exc)},
        )
        return None, "Dokument konnte nicht geoeffnet werden"
    finally:
        try:
            if dokument is not None:
                dokument.Close(False)
        except Exception:
            pass
        try:
            if app is not None:
                app.Quit()
        except Exception:
            pass
        if temp_pfad:
            try:
                Path(temp_pfad).unlink(missing_ok=True)
            except Exception:
                pass
        if com_aktiv:
            try:
                pythoncom.CoUninitialize()  # type: ignore[name-defined]
            except Exception:
                pass


def _oeffne_dokument(filepath: str, erlaubte_ordner: list[Path], read_only: bool = False):
    try:
        datei = Path(filepath)
        if not datei.exists() or datei.suffix.lower() != ".docx":
            _debug_word_ereignis(
                "word.open.validation",
                "invalid_path_or_ext",
                filepath,
                f"exists={datei.exists()};suffix={datei.suffix.lower()}",
            )
            return None, "Datei nicht gefunden oder kein .docx"
        if not _ist_in_erlaubten_ordnern(datei, erlaubte_ordner):
            _debug_word_ereignis(
                "word.open.validation",
                "outside_allowed_folders",
                filepath,
                f"allowed_count={len(erlaubte_ordner)}",
            )
            return None, "Datei liegt ausserhalb der erlaubten Ordner"
        dokument, fehler = _oeffne_docx_mit_retry(datei)
        if not fehler:
            return dokument, ""
        if read_only:
            live_dokument, live_fehler = _oeffne_docx_ueber_laufendes_word(datei)
            if not live_fehler:
                return live_dokument, ""
            getobj_dokument, getobj_fehler = _oeffne_docx_ueber_word_getobject(datei)
            if not getobj_fehler:
                return getobj_dokument, ""
            fallback_fehler = "Dokument konnte nicht geoeffnet werden"
            if _temp_copy_fallback_aktiv():
                fallback_dokument, fallback_fehler = _oeffne_docx_ueber_tempkopie(datei)
                if not fallback_fehler:
                    return fallback_dokument, ""
            com_dokument, com_fehler = _oeffne_docx_ueber_com_readonly(datei)
            if not com_fehler:
                return com_dokument, ""
            _debug_word_ereignis(
                "word.open.failed",
                "open_and_fallback_failed",
                filepath,
                "direct_and_fallback_failed",
                {
                    "fallback_used": True,
                    "mode": "read",
                    "live_error": str(live_fehler or "")[:80],
                    "getobj_error": str(getobj_fehler or "")[:80],
                    "temp_error": str(fallback_fehler or "")[:80],
                    "com_error": str(com_fehler or "")[:80],
                },
            )
            return None, com_fehler or fallback_fehler or getobj_fehler or live_fehler
        return None, fehler
    except Exception:
        _debug_word_ereignis("word.open.exception", "open_exception", filepath)
        return None, "Dokument konnte nicht geoeffnet werden"


def _tabellen_als_liste(dokument) -> list:
    tabellen = []
    for tabelle in dokument.tables:
        zeilen = []
        for zeile in tabelle.rows:
            zeilen.append([zelle.text.strip() for zelle in zeile.cells])
        tabellen.append(zeilen)
    return tabellen


def _normalisiere_uhrzeit_text(text: str) -> str:
    roh = str(text or "").strip().lower().replace("uhr", "").replace(".", ":")
    kopf = roh.split("\n", 1)[0].strip()
    treffer = re.search(r"(\d{1,2})(?::(\d{1,2}))?", kopf)
    if treffer is None:
        return ""
    stunde = int(treffer.group(1))
    minute = int(treffer.group(2) or "0")
    if stunde > 23 or minute > 59:
        return ""
    return f"{stunde:02d}:{minute:02d}"


def _finde_zeile_fuer_uhrzeit(tabelle, uhrzeit: str):
    zielzeit = _normalisiere_uhrzeit_text(uhrzeit)
    if not zielzeit:
        return None
    for zeile in tabelle.rows[1:]:
        zellen = zeile.cells
        if len(zellen) <= 1:
            continue
        feldzeit = _normalisiere_uhrzeit_text(zellen[0].text)
        if feldzeit == zielzeit:
            return zeile
    return None


def _speichere_dokument(dokument, filepath: str) -> str:
    letzter_fehler = "Schreibvorgang fehlgeschlagen"
    for versuch in range(SPEICHERN_RETRY):
        try:
            dokument.save(filepath)
            time.sleep(NACH_SPEICHERN_WARTEZEIT)
            return ""
        except PermissionError:
            letzter_fehler = "Datei ist gesperrt (z. B. in Word geoeffnet)"
        except OSError:
            letzter_fehler = "Datei konnte nicht gespeichert werden"
        except Exception:
            letzter_fehler = "Schreibvorgang fehlgeschlagen"
        if versuch < SPEICHERN_RETRY - 1:
            time.sleep(SPEICHERN_WARTEZEIT)
    return letzter_fehler


def list_files_in_folder() -> dict:
    try:
        ordner_liste, fehler = _hole_erlaubte_ordner()
        if fehler:
            return {"fehler": fehler}
        dateien = set()
        for ordner in ordner_liste:
            for datei in ordner.glob("*.docx"):
                dateien.add(datei.name)
        return {"dateien": sorted(dateien)}
    except Exception:
        return {"fehler": "Dateiliste konnte nicht gelesen werden"}


def search_docx_files(dateiname: str = "", datum: str = "") -> dict:
    try:
        ordner_liste, fehler = _hole_erlaubte_ordner()
        if fehler:
            return {"fehler": fehler}
        suchname = str(dateiname or "").strip().lower()
        suchdatum = _normalisiere_datum(datum) if datum else ""
        datum_varianten = _datum_varianten(suchdatum)
        if datum and not suchdatum:
            return {"fehler": "Datum hat ein ungueltiges Format"}
        ergebnisse = []
        for eintrag in _hole_docx_index(ordner_liste):
            dateiname_feld = str(eintrag.get("dateiname", "") or "")
            datei_datum = str(eintrag.get("datum", "") or "")
            if suchname and suchname not in dateiname_feld.lower():
                continue
            if suchdatum and datei_datum not in datum_varianten:
                continue
            ergebnisse.append({
                "dateiname": dateiname_feld,
                "pfad": str(eintrag.get("pfad", "") or ""),
                "datum": datei_datum,
            })
        return {"dateien": ergebnisse}
    except Exception:
        return {"fehler": "Dateisuche fehlgeschlagen"}


def read_word_file(filepath: str) -> dict:
    ordner_liste, fehler = _hole_erlaubte_ordner()
    if fehler:
        return {"fehler": fehler}
    dokument, fehler = _oeffne_dokument(filepath, ordner_liste, read_only=True)
    if fehler:
        return {"fehler": fehler}
    try:
        freitext = "\n".join(absatz.text for absatz in dokument.paragraphs)
        tabellen = _tabellen_als_liste(dokument)
        return {"freitext": freitext, "tabellen": tabellen}
    except Exception:
        return {"fehler": "Datei konnte nicht gelesen werden"}


def find_free_slots(filepath: str) -> dict:
    ordner_liste, fehler = _hole_erlaubte_ordner()
    if fehler:
        return {"fehler": fehler}
    dokument, fehler = _oeffne_dokument(filepath, ordner_liste, read_only=True)
    if fehler:
        return {"fehler": fehler}
    try:
        if not dokument.tables:
            return {"fehler": "Keine Tabelle in der Datei gefunden"}
        freie_slots = []
        for zeile in dokument.tables[0].rows[1:]:
            zellen = zeile.cells
            if len(zellen) > 1 and not zellen[1].text.strip():
                freie_slots.append(zellen[0].text.strip())
        return {"freie_slots": freie_slots, "datei": Path(filepath).name}
    except Exception:
        return {"fehler": "Freie Slots konnten nicht ermittelt werden"}


def find_file_by_date(datum: str) -> dict:
    try:
        normalisiert = _normalisiere_datum(datum)
        if not normalisiert:
            _debug_word_ereignis("word.find_by_date", "invalid_date", detail="normalisierung_leer")
            return {"fehler": "Datum hat ein ungueltiges Format"}
        ergebnis = search_docx_files(datum=normalisiert)
        if "fehler" in ergebnis:
            _debug_word_ereignis("word.find_by_date", "search_failed", detail="search_docx_files_fehler")
            return ergebnis
        dateien = ergebnis.get("dateien", [])
        if not dateien:
            _debug_word_ereignis("word.find_by_date", "no_file_for_date", detail=f"datum={normalisiert}")
            return {"fehler": "Keine Datei fuer dieses Datum gefunden"}
        return {"filepath": str(dateien[0].get("pfad", ""))}
    except Exception:
        _debug_word_ereignis("word.find_by_date", "find_by_date_exception")
        return {"fehler": "Dateisuche fehlgeschlagen"}


def _text_normalisieren(text: str) -> str:
    roh = str(text or "").strip().lower()
    roh = unicodedata.normalize("NFKD", roh)
    roh = "".join(ch for ch in roh if not unicodedata.combining(ch))
    return " ".join(roh.split())


def _eintrag_aus_zeile(zeile, filepath: str, datei_datum: str) -> dict:
    if len(zeile.cells) < 2:
        return {}
    uhrzeit_raw = str(zeile.cells[0].text or "").strip()
    name = str(zeile.cells[1].text or "").strip()
    if not name:
        return {}
    uhrzeit = _normalisiere_uhrzeit_text(uhrzeit_raw) or uhrzeit_raw
    return {
        "datum": datei_datum,
        "dateiname": Path(filepath).name,
        "filepath": filepath,
        "uhrzeit": uhrzeit,
        "name": name,
    }


def _zeile_passt_auf_suche(zeile, suchtext: str, zielzeit: str) -> bool:
    zeit = _normalisiere_uhrzeit_text(zeile.cells[0].text if len(zeile.cells) > 0 else "")
    if zielzeit and zeit != zielzeit:
        return False
    if not suchtext:
        return True
    gesamt = " ".join(str(zelle.text or "") for zelle in zeile.cells)
    return suchtext in _text_normalisieren(gesamt)


def _fuege_treffer_aus_dokument_hinzu(dokument, dateipfad: str, datei_datum: str, such_norm: str, zielzeit: str, treffer: list) -> None:
    for zeile in dokument.tables[0].rows[1:]:
        if not _zeile_passt_auf_suche(zeile, such_norm, zielzeit):
            continue
        daten = _eintrag_aus_zeile(zeile, dateipfad, datei_datum)
        if daten:
            treffer.append(daten)


def _lese_dokument_fuer_suche(dateipfad: str, ordner_liste: list[Path]):
    dokument, lesefehler = _oeffne_dokument(dateipfad, ordner_liste, read_only=True)
    if lesefehler:
        _debug_word_ereignis(
            "word.search.skip",
            "search_open_failed",
            dateipfad,
            f"open_failed:{lesefehler}",
            {"mode": "read"},
        )
        return None, "open_failed"
    if not dokument.tables:
        _debug_word_ereignis("word.search.skip", "search_no_table", dateipfad)
        return None, "no_table"
    return dokument, ""


def _retry_gescheiterte_dateien(
    fehlerpfade: list[str],
    ordner_liste: list[Path],
    treffer: list,
    such_norm: str,
    zielzeit: str,
    daten_map: dict,
    stats: dict,
) -> int:
    if not fehlerpfade:
        return 0
    time.sleep(0.25)
    noch_fehler = 0
    for dateipfad in fehlerpfade:
        if _scan_limit_erreicht(stats):
            break
        dokument, lesefehler = _lese_dokument_fuer_suche(dateipfad, ordner_liste)
        if lesefehler:
            noch_fehler += 1
            stats["retry_skipped"] = int(stats.get("retry_skipped", 0)) + 1
            stats["skipped"] = int(stats.get("skipped", 0)) + 1
            _scan_limit_erreicht(stats)
            continue
        stats["retry_opened"] = int(stats.get("retry_opened", 0)) + 1
        stats["opened"] = int(stats.get("opened", 0)) + 1
        datei_datum = str(daten_map.get(dateipfad, "") or "")
        _fuege_treffer_aus_dokument_hinzu(dokument, dateipfad, datei_datum, such_norm, zielzeit, treffer)
    return noch_fehler


def _timing_ms(start_ts: float) -> int:
    dauer = (time.time() - float(start_ts or 0.0)) * 1000.0
    return max(0, int(dauer))


def _search_mit_index_oder_scan(
    datei_suche: dict,
    ordner_liste: list[Path],
    such_norm: str,
    zielzeit: str,
    suchdatum: str,
) -> tuple[dict, str, dict]:
    gesamt_stats = {
        "opened": 0,
        "skipped": 0,
        "retry_opened": 0,
        "retry_skipped": 0,
        "limit_hit": False,
        "limit_reason": "",
        "max_open_attempts": 0,
        "max_skip": 0,
        "stages": "",
    }

    def _merge_stats(stats: dict, stage_name: str) -> None:
        if not isinstance(stats, dict):
            return
        for key in ["opened", "skipped", "retry_opened", "retry_skipped"]:
            gesamt_stats[key] = int(gesamt_stats.get(key, 0)) + int(stats.get(key, 0) or 0)
        if bool(stats.get("limit_hit")):
            gesamt_stats["limit_hit"] = True
            gesamt_stats["limit_reason"] = str(stats.get("limit_reason", "") or "")
        if int(stats.get("max_open_attempts", 0) or 0) > 0:
            gesamt_stats["max_open_attempts"] = int(stats.get("max_open_attempts", 0) or 0)
        if int(stats.get("max_skip", 0) or 0) > 0:
            gesamt_stats["max_skip"] = int(stats.get("max_skip", 0) or 0)
        stages = [s for s in str(gesamt_stats.get("stages", "")).split(",") if s]
        if stage_name and stage_name not in stages:
            stages.append(stage_name)
        gesamt_stats["stages"] = ",".join(stages)

    if not _index_feature_aktiv():
        limits = _scan_guardrail_limits("limited")
        scan = _search_appointments_scan(
            datei_suche,
            ordner_liste,
            such_norm,
            zielzeit,
            limits["max_open_attempts"],
            limits["max_skip"],
            "scan_limited",
        )
        _merge_stats(scan.get("scan_stats", {}), "scan_limited")
        return scan, "scan", gesamt_stats

    index_ergebnis = _search_appointments_index(datei_suche, ordner_liste, such_norm, zielzeit)
    if "fehler" not in index_ergebnis and index_ergebnis.get("treffer"):
        _merge_stats(index_ergebnis.get("scan_stats", {}), "index")
        return index_ergebnis, "index", gesamt_stats

    limits = _scan_guardrail_limits("limited")
    limited = _search_appointments_scan(
        datei_suche,
        ordner_liste,
        such_norm,
        zielzeit,
        limits["max_open_attempts"],
        limits["max_skip"],
        "scan_limited",
    )
    _merge_stats(limited.get("scan_stats", {}), "scan_limited")

    if limited.get("treffer"):
        if "warnung" not in limited and "fehler" in index_ergebnis:
            limited["warnung"] = "Index nicht verfuegbar, Suche ueber begrenzten Scan"
        return limited, "scan_limited", gesamt_stats

    if limited.get("erweiterter_vollscan_moeglich") and not (suchdatum or zielzeit):
        if "warnung" not in limited:
            limited["warnung"] = "Suche frueh beendet. Mit Datum oder Uhrzeit eingrenzen fuer Vollscan."
        return limited, "scan_limited_guarded", gesamt_stats

    full_limits = _scan_guardrail_limits("full")
    full = _search_appointments_scan(
        datei_suche,
        ordner_liste,
        such_norm,
        zielzeit,
        full_limits["max_open_attempts"],
        full_limits["max_skip"],
        "scan_full",
    )
    _merge_stats(full.get("scan_stats", {}), "scan_full")
    if "warnung" not in full and "fehler" in index_ergebnis:
        full["warnung"] = "Index nicht verfuegbar, Suche ueber Vollscan"
    return full, "scan_fallback", gesamt_stats


def search_appointments(suchtext: str = "", uhrzeit: str = "", datum: str = "") -> dict:
    start_ts = time.time()
    ordner_liste, fehler = _hole_erlaubte_ordner()
    if fehler:
        return {"fehler": fehler}
    suchdatum = _normalisiere_datum(datum) if datum else ""
    if datum and not suchdatum:
        return {"fehler": "Datum hat ein ungueltiges Format"}
    datei_suche = search_docx_files(datum=suchdatum) if suchdatum else search_docx_files()
    if "fehler" in datei_suche:
        return datei_suche
    such_norm = _text_normalisieren(suchtext)
    zielzeit = _normalisiere_uhrzeit_text(uhrzeit) if uhrzeit else ""
    ergebnis, quelle, stats = _search_mit_index_oder_scan(datei_suche, ordner_liste, such_norm, zielzeit, suchdatum)
    dauer_ms = _timing_ms(start_ts)
    treffer = ergebnis.get("treffer", []) if isinstance(ergebnis, dict) else []
    dateien = datei_suche.get("dateien", []) if isinstance(datei_suche, dict) else []
    _logge_such_performance_erweitert(str(quelle or "scan"), dauer_ms, len(treffer), len(dateien), stats)
    return ergebnis if isinstance(ergebnis, dict) else {"fehler": "Dateisuche fehlgeschlagen"}


def _vorbereite_schreibvorgang(filepath: str, bestaetigt: bool):
    if not bestaetigt:
        return None, "", "Bestaetigung erforderlich"
    ordner_liste, fehler = _hole_erlaubte_ordner()
    if fehler:
        return None, "", fehler
    dokument, fehler = _oeffne_dokument(filepath, ordner_liste, read_only=False)
    if fehler:
        return None, "", fehler
    backup_pfad = create_backup(filepath)
    if not backup_pfad:
        return None, "", "Backup konnte nicht erstellt werden"
    return dokument, backup_pfad, ""


def _erster_run_mit_format(zelle):
    for absatz in zelle.paragraphs:
        for run in absatz.runs:
            if str(run.text or "").strip():
                return run
    for absatz in zelle.paragraphs:
        if absatz.runs:
            return absatz.runs[0]
    return None


def _kopiere_run_format(referenz_run, ziel_run) -> None:
    if referenz_run is None:
        return
    try:
        ziel_run.style = referenz_run.style
    except Exception:
        pass
    ziel_run.font.name = referenz_run.font.name
    ziel_run.font.size = referenz_run.font.size
    ziel_run.font.bold = referenz_run.font.bold
    ziel_run.font.italic = referenz_run.font.italic
    ziel_run.font.underline = referenz_run.font.underline
    ziel_run.font.color.rgb = referenz_run.font.color.rgb


def _setze_zellentext_formatiert(zelle, text: str) -> None:
    referenz_run = _erster_run_mit_format(zelle)
    absatz = zelle.paragraphs[0] if zelle.paragraphs else zelle.add_paragraph()
    try:
        absatz.clear()
    except Exception:
        absatz.text = ""
    run = absatz.add_run(str(text or "").strip())
    _kopiere_run_format(referenz_run, run)


def _slot_ist_belegt(zeile) -> bool:
    return bool(str(zeile.cells[1].text or "").strip())


def _lese_terminwerte_aus_zeile(zeile) -> dict:
    return {
        "name": str(zeile.cells[1].text or "").strip(),
        "svnr": str(zeile.cells[2].text or "").strip(),
        "geburtsdatum": str(zeile.cells[3].text or "").strip(),
        "adresse": str(zeile.cells[4].text or "").strip(),
        "firma": str(zeile.cells[5].text or "").strip(),
        "untersuchungsart": str(zeile.cells[6].text or "").strip(),
        "vgue": str(zeile.cells[8].text or "").strip(),
    }


def _merge_terminwerte(neu: dict, alt: dict, ersetzen: bool) -> dict:
    if not ersetzen:
        return dict(neu)
    gemergt = {}
    for feld, neuer_wert in neu.items():
        text = str(neuer_wert or "").strip()
        if text:
            gemergt[feld] = text
            continue
        gemergt[feld] = str(alt.get(feld, "") or "").strip()
    return gemergt


def _trage_terminwerte_ein(zeile, werte: dict) -> None:
    name = str(werte.get("name", "") or "")
    svnr = str(werte.get("svnr", "") or "")
    geburtsdatum = str(werte.get("geburtsdatum", "") or "")
    adresse = str(werte.get("adresse", "") or "")
    firma = str(werte.get("firma", "") or "")
    art = str(werte.get("untersuchungsart", "") or "")
    vgue = str(werte.get("vgue", "") or "")
    werte = [(1, name), (2, svnr), (3, geburtsdatum), (4, adresse), (5, firma), (6, art), (8, vgue)]
    for index, wert in werte:
        _setze_zellentext_formatiert(zeile.cells[index], wert)


def write_appointment(
    filepath: str,
    uhrzeit: str,
    name: str,
    svnr: str = "",
    geburtsdatum: str = "",
    adresse: str = "",
    firma: str = "",
    untersuchungsart: str = "",
    vgue: str = "",
    ersetzen: bool = False,
    bestaetigt: bool = False,
) -> dict:
    expired_hinweis = _pending_expired_hinweis(filepath, uhrzeit)
    if expired_hinweis:
        return _fehler_result(expired_hinweis, "expired")
    neue_werte = {
        "name": str(name or "").strip(),
        "svnr": str(svnr or "").strip(),
        "geburtsdatum": str(geburtsdatum or "").strip(),
        "adresse": str(adresse or "").strip(),
        "firma": str(firma or "").strip(),
        "untersuchungsart": str(untersuchungsart or "").strip(),
        "vgue": str(vgue or "").strip(),
    }
    dokument, backup_pfad, fehler = _vorbereite_schreibvorgang(filepath, bestaetigt)
    if fehler:
        if _ist_lock_fehler(fehler):
            _debug_word_ereignis("word.lock.detected", "write_prepare_lock", filepath)
            return _pending_enqueue(filepath, uhrzeit, neue_werte, ersetzen, "")
        return _fehler_result(fehler, "prepare_failed")
    try:
        if not dokument.tables:
            return _fehler_result("Keine Tabelle in der Datei gefunden", "no_table")
        zeile = _finde_zeile_fuer_uhrzeit(dokument.tables[0], uhrzeit)
        if zeile is None:
            return _fehler_result("Slot nicht gefunden", "slot_not_found")
        ist_belegt = _slot_ist_belegt(zeile)
        if ist_belegt and not ersetzen:
            return _fehler_result("Slot bereits belegt", "slot_belegt")
        alte_werte = _lese_terminwerte_aus_zeile(zeile) if ist_belegt else {}
        zielwerte = _merge_terminwerte(neue_werte, alte_werte, ersetzen)
        if not str(zielwerte.get("name", "") or "").strip():
            return _fehler_result("Name fehlt", "name_required")
        _trage_terminwerte_ein(zeile, zielwerte)
        speicher_fehler = _speichere_dokument(dokument, filepath)
        if speicher_fehler:
            if _ist_lock_fehler(speicher_fehler):
                _debug_word_ereignis("word.lock.detected", "write_save_lock", filepath)
                return _pending_enqueue(filepath, uhrzeit, neue_werte, ersetzen, backup_pfad)
            return _fehler_result(speicher_fehler, "save_failed")
        return {"erfolg": True, "backup_pfad": backup_pfad}
    except Exception:
        return _fehler_result("Schreibvorgang fehlgeschlagen", "write_exception")

